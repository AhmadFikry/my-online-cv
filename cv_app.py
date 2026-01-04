import streamlit as st
import os
import re
from PyPDF2 import PdfReader
from crewai import Agent, Task, Crew, LLM
from crewai_tools import ScrapeWebsiteTool, SerperDevTool, FileReadTool
from docx import Document
import io
from dotenv import load_dotenv

# --- 0. LOAD LOCAL KEYS ---
load_dotenv() 

# --- 1. PAGE CONFIG ---
st.set_page_config(page_title="HR CV Architect Pro", page_icon="⚖️", layout="wide")

# --- 2. HELPERS ---
def extract_pdf_text(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        content = page.extract_text()
        if content:
            text += content
    return text

def clean_markdown(text):
    text = re.sub(r'#+\s*', '', text) 
    text = text.replace('**', '').replace('__', '')
    return text

def create_docx(text):
    doc = Document()
    doc.add_heading('Tailored HR Resume & Strategy', 0)
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.isupper() and len(line) < 60:
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. UTILITY SIDEBAR ---
with st.sidebar:
    st.title("⚙️ Workspace")
    if st.button("🗑️ Reset Application"):
        st.session_state.clear()
        st.rerun()
    st.divider()
    st.caption("CV Architect v1.6 - HR Edition")

# --- 4. MAIN UI ---
st.title("⚖️ HR CV Architect Pro")
st.markdown("##### Strategic HR Resume Engineering & Behavioral Interview Prep")

st.subheader("STEP 1: YOUR DATA")
col_cv, col_links = st.columns(2)
with col_cv:
    uploaded_file = st.file_uploader("Upload Current Resume (PDF/TXT)", type=["pdf", "txt"])
with col_links:
    linkedin_url = st.text_input("LinkedIn/Portfolio URL", placeholder="https://linkedin.com/in/your-profile")

st.subheader("STEP 2: THE TARGET")
job_url = st.text_input("Job Posting URL", placeholder="Paste the HR job link here")
personal_writeup = st.text_area("Key HR Achievements & Values", 
                                placeholder="Describe your experience with DEI, Labor Law, or Retention...", height=120)

if 'ai_result' not in st.session_state:
    st.session_state.ai_result = ""
if 'interview_prep' not in st.session_state:
    st.session_state.interview_prep = ""

# --- 5. EXECUTION ---
st.write("---")
if st.button("🚀 GENERATE HR APPLICATION PACKAGE", use_container_width=True):
    gemini_key = os.environ.get("GEMINI_API_KEY")
    groq_key = os.environ.get("GROQ_API_KEY")

    if not gemini_key or not groq_key:
        st.error("⚠️ Connection Error: Missing API Keys in .env file.")
    elif not uploaded_file or not job_url:
        st.warning("⚠️ Missing inputs: Please provide at least a Resume and Job URL.")
    else:
        # Extract text once to avoid tool errors
        cv_text = extract_pdf_text(uploaded_file) if uploaded_file.type == "application/pdf" else uploaded_file.read().decode()

        # --- HYBRID LLM INITIALIZATION ---
        research_llm = LLM(
            model="gemini/gemini-2.5-flash", 
            api_key=gemini_key,
            temperature=0.0
        )

        logic_llm = LLM(
           model="groq/llama-3.3-70b-versatile", 
           api_key=groq_key,
           temperature=0.0
        )

        # --- TOOLS (REDUCED TO PREVENT HANGING) ---
        search_tool = SerperDevTool()
        read_resume = FileReadTool()
        # Note: We are NOT using ScrapeWebsiteTool here because it often hangs on LinkedIn

# --- HR-SPECIALIZED AGENTS (Fixed with required Backstories) ---
        researcher = Agent(
            role="HR Market Analyst",
            goal="Identify core HR competencies and requirements from the job link provided.",
            backstory="You are an expert in Talent Acquisition and Labor Law with 15 years of experience in market analysis.",
            tools=[search_tool], 
            verbose=True,
            llm=research_llm,
            max_iter=2,
            allow_delegation=False
        )

        profiler = Agent(
            role="HR Professional Profiler",
            goal="Dissect background to highlight leadership and strategic impact.",
            backstory="You are a specialist in Executive Coaching and HR Career Development, expert at spotting high-potential traits.",
            tools=[], 
            verbose=True,
            llm=research_llm,
            max_iter=2,
            allow_delegation=False
        )

        resume_strategist = Agent(
            role="HR Executive Resume Strategist",
            goal="Tailor the resume to pass ATS and appeal to senior HR leadership.",
            backstory="You are a veteran HR Director who has reviewed thousands of resumes and knows exactly what triggers a 'Yes' for an interview.",
            tools=[], 
            verbose=True,
            llm=logic_llm,
            max_iter=2,
            allow_delegation=False
        )

        interview_preparer = Agent(
            role="HR Interview Coach",
            goal="Formulate behavioral and situational HR interview questions.",
            backstory="You are an expert interviewer trained in the STAR method and organizational psychology.",
            tools=[],
            verbose=True,
            llm=logic_llm,
            max_iter=2,
            allow_delegation=False
        )

        # --- HR-SPECIFIC TASKS ---
        research_task = Task(
            description=f"Review this Job Link: {job_url}. If you cannot access it, focus on the standard requirements for an HR role with that title.",
            expected_output="Dossier of HR requirements.",
            agent=researcher
        )

        profile_task = Task(
            description=f"Analyze this candidate text: {cv_text}. And this context: {personal_writeup}.",
            expected_output="Profile of the candidate's HR leadership style.",
            agent=profiler
        )

        resume_strategy_task = Task(
            description=(
                "Tailor the HR resume using ONLY the information provided in the profile and CV. "
                "CRITICAL: Do not invent certifications, degrees, or job history. "
                "If a certificate is listed as 'Completed' in the CV, do not change it to 'Pursuing'. "
                "Focus on re-wording existing achievements to match the job's keywords rather than adding new ones."
            ),
            expected_output="A factually accurate but strategically worded HR resume.",
            context=[research_task, profile_task],
            agent=resume_strategist
        )

        interview_preparation_task = Task(
            description="Generate 5 situational and 5 strategic interview questions.",
            expected_output="HR interview prep guide.",
            context=[research_task, profile_task, resume_strategy_task],
            agent=interview_preparer
        )

        crew = Crew(
            agents=[researcher, profiler, resume_strategist, interview_preparer], 
            tasks=[research_task, profile_task, resume_strategy_task, interview_preparation_task],
            process=st.session_state.get('process', 'sequential') # Sequential is safer for local debugging
        )

        with st.status("⚖️ Strategizing HR Career Move...", expanded=True) as status:
            # We use a standard kickoff here to ensure Streamlit catches the return
            result = crew.kickoff()
            
            # Capture the specific task outputs
            st.session_state.ai_result = str(resume_strategy_task.output.raw)
            st.session_state.interview_prep = str(interview_preparation_task.output.raw)
            status.update(label="✅ Success!", state="complete")
        
        st.rerun()

# --- 6. RESULTS ---
if st.session_state.ai_result:
    t_cv, t_int = st.tabs(["📄 New HR Resume", "🎙️ HR Interview Prep"])
    with t_cv:
        st.markdown(st.session_state.ai_result)
        st.download_button("📥 Word (.docx)", data=create_docx(clean_markdown(st.session_state.ai_result)), file_name="HR_Resume.docx")
    with t_int:
        st.markdown(st.session_state.interview_prep)